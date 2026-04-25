from docx import Document
from docx.shared import Inches
from src.safwanbuddy.core import logger
import os

class WordGenerator:
    def __init__(self):
        self.output_dir = "output/documents"
        os.makedirs(self.output_dir, exist_ok=True)

    def create_document(self, title: str, content: list, filename: str = "document.docx"):
        """
        Creates a Word document.
        content is a list of dicts: [{"type": "paragraph", "text": "..."}, {"type": "heading", "level": 1, "text": "..."}]
        """
        try:
            doc = Document()
            doc.add_heading(title, 0)

            for item in content:
                c_type = item.get("type")
                text = item.get("text", "")
                
                if c_type == "paragraph":
                    doc.add_paragraph(text)
                elif c_type == "heading":
                    doc.add_heading(text, level=item.get("level", 1))
                elif c_type == "bullet":
                    doc.add_paragraph(text, style='List Bullet')
                elif c_type == "numbered":
                    doc.add_paragraph(text, style='List Number')

            path = os.path.join(self.output_dir, filename)
            doc.save(path)
            logger.info(f"Word document saved to {path}")
            return path
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            return None

word_generator = WordGenerator()
