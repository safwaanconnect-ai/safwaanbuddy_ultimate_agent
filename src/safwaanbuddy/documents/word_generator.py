"""Word document generation using python-docx."""

import logging
from pathlib import Path
from typing import List, Dict, Any, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available")

from ..core.events import EventBus, EventType
from ..core.config import ConfigManager


class WordGenerator:
    """Generate Word documents."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.config = ConfigManager()
        self.event_bus = EventBus()
        
        self.document: Optional[Document] = None
    
    def create_document(self) -> bool:
        """Create new document.
        
        Returns:
            True if successful
        """
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available")
            return False
        
        try:
            self.document = Document()
            self.logger.info("Created new Word document")
            return True
        except Exception as e:
            self.logger.error(f"Failed to create document: {e}")
            return False
    
    def add_heading(self, text: str, level: int = 1) -> bool:
        """Add heading to document.
        
        Args:
            text: Heading text
            level: Heading level (1-9)
            
        Returns:
            True if successful
        """
        if not self.document:
            return False
        
        try:
            self.document.add_heading(text, level=level)
            return True
        except Exception as e:
            self.logger.error(f"Failed to add heading: {e}")
            return False
    
    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False) -> bool:
        """Add paragraph to document.
        
        Args:
            text: Paragraph text
            bold: Bold text
            italic: Italic text
            
        Returns:
            True if successful
        """
        if not self.document:
            return False
        
        try:
            paragraph = self.document.add_paragraph(text)
            if bold:
                paragraph.runs[0].bold = True
            if italic:
                paragraph.runs[0].italic = True
            return True
        except Exception as e:
            self.logger.error(f"Failed to add paragraph: {e}")
            return False
    
    def add_table(self, data: List[List[str]], has_header: bool = True) -> bool:
        """Add table to document.
        
        Args:
            data: Table data (list of rows)
            has_header: First row is header
            
        Returns:
            True if successful
        """
        if not self.document or not data:
            return False
        
        try:
            rows = len(data)
            cols = len(data[0])
            
            table = self.document.add_table(rows=rows, cols=cols)
            table.style = 'Light Grid Accent 1'
            
            for i, row in enumerate(data):
                for j, cell_text in enumerate(row):
                    table.rows[i].cells[j].text = str(cell_text)
                    if has_header and i == 0:
                        table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
            
            return True
        except Exception as e:
            self.logger.error(f"Failed to add table: {e}")
            return False
    
    def save_document(self, filepath: Path) -> bool:
        """Save document to file.
        
        Args:
            filepath: Output file path
            
        Returns:
            True if successful
        """
        if not self.document:
            return False
        
        try:
            filepath = Path(filepath)
            filepath.parent.mkdir(parents=True, exist_ok=True)
            
            self.document.save(filepath)
            
            self.event_bus.emit(EventType.DOCUMENT_SAVED, {
                "type": "word",
                "filepath": str(filepath)
            })
            
            self.logger.info(f"Document saved: {filepath}")
            return True
        except Exception as e:
            self.logger.error(f"Failed to save document: {e}")
            return False
